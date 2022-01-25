from sanic import Sanic
from sanic_cors import CORS
from sanic.response import json
from docx import Document
from io import BytesIO
import base64
<<<<<<< HEAD
app = Sanic("LDA_api")
=======
app = Sanic("My Hello, world app")
>>>>>>> 757df8018a2b0c3c63115b48c96929840a51d3a2
CORS(app)

# define model here
# model = tf.keras.models.load_model('path/to/model/h5py')

@app.route('/')
async def test(request):
    return json({'hello': 'world'})

<<<<<<< HEAD
@app.route('/txt', methods=['POST'])
async def txt_reader(request):
    json_body = request.json
    file = json_body['file']
    file = base64.b64decode(file)
    file = BytesIO(file).read().decode('utf-8')
    return json({'response': file})

@app.route('/doc', methods=['POST'])
@app.route('/docx', methods=['POST'])
async def docx_reader(request):
=======
@app.route('/doc', methods=['POST'])
@app.route('/docx', methods=['POST'])
async def test(request):
>>>>>>> 757df8018a2b0c3c63115b48c96929840a51d3a2
    json_body = request.json
    file = json_body['file']
    file = base64.b64decode(file)
    doc = Document(BytesIO(file))
    res = {}
<<<<<<< HEAD
    paragraphs = [p.text for p in doc.paragraphs]
    for i, p in enumerate(paragraphs):
        res[f'p{i}'] = p
=======
    for i, para in enumerate(doc.paragraphs):
        res[str(i)] = para.text
>>>>>>> 757df8018a2b0c3c63115b48c96929840a51d3a2
    return json(res)


host = '127.0.0.1'
port = 8000
if __name__ == '__main__':
    app.run(host=host, port=port, auto_reload=True)